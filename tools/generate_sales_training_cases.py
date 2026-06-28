"""批量生成销售训练知识库示例文档。

这个脚本只生成本地 docx、csv 和说明文档，不会连接数据库、MinIO 或向量库。
生成出的 docx 可用于手动上传到销售训练资料库，测试切片、预览、发布和销售陪练效果。
"""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "generated_knowledge" / "sales_training_cases"
DOC_COUNT = 50


INDUSTRIES = [
    ("机械配件加工", "非标农机连接件、工程机械轴套、金属压铸件", "山东潍坊", "L1一星启航版 ¥11800/年"),
    ("农机设备", "小型耕地机、除草机、灌溉泵、田园转运车", "河南洛阳", "标准版 2.98 万元"),
    ("锂电叉车", "锂电座驾式叉车、仓储堆高车、电动搬运车", "山东德州", "L5五星旗舰版 ¥199800/年"),
    ("家居灯具", "吊灯、壁灯、庭院灯、酒店工程灯", "广东中山", "L2二星扬帆版 ¥39800/年"),
    ("包装机械", "封箱机、贴标机、热收缩包装线、灌装设备", "浙江温州", "L3三星远航版 ¥69800/年"),
    ("汽摩配件", "刹车片、滤清器、减震器、改装件", "河北邢台", "L2二星扬帆版 ¥39800/年"),
    ("五金工具", "扳手、套筒、钳类工具、工具箱", "浙江金华", "L3三星远航版 ¥69800/年"),
    ("宠物用品", "宠物牵引绳、猫爬架、宠物窝、喂食器", "江苏苏州", "L2二星扬帆版 ¥39800/年"),
    ("户外家具", "藤编桌椅、庭院伞、露营桌椅、遮阳棚", "浙江宁波", "L4四星领航版 ¥128000/年"),
    ("工业泵阀", "离心泵、隔膜泵、球阀、蝶阀、管件", "江苏盐城", "L4四星领航版 ¥128000/年"),
    ("建材卫浴", "花洒、龙头、陶瓷盆、浴室柜", "福建泉州", "L3三星远航版 ¥69800/年"),
    ("安防设备", "摄像头、门禁、报警器、NVR 设备", "广东深圳", "L4四星领航版 ¥128000/年"),
    ("新能源配件", "光伏支架、储能外壳、充电桩配件", "江苏常州", "L5五星旗舰版 ¥199800/年"),
    ("塑料制品", "收纳箱、周转筐、注塑外壳、包装盒", "浙江台州", "L2二星扬帆版 ¥39800/年"),
    ("食品加工设备", "切菜机、和面机、真空包装机、烘干线", "山东诸城", "L3三星远航版 ¥69800/年"),
    ("医疗耗材", "一次性手套、敷料、采样管、护理耗材", "湖北仙桃", "L4四星领航版 ¥128000/年"),
    ("纺织面料", "功能面料、户外布料、家纺面料、针织布", "江苏南通", "L3三星远航版 ¥69800/年"),
    ("厨房用品", "不锈钢锅具、刀具、硅胶厨具、收纳架", "广东阳江", "L2二星扬帆版 ¥39800/年"),
    ("儿童玩具", "益智玩具、积木、遥控车、户外玩具", "浙江义乌", "L3三星远航版 ¥69800/年"),
    ("办公家具", "升降桌、办公椅、文件柜、会议桌", "广东佛山", "L4四星领航版 ¥128000/年"),
    ("服装辅料", "拉链、纽扣、织带、吊牌、包装袋", "浙江嘉兴", "L2二星扬帆版 ¥39800/年"),
    ("电子元器件", "连接器、线束、传感器、电源模块", "广东东莞", "L5五星旗舰版 ¥199800/年"),
    ("园林工具", "电链锯、割草机、吹风机、修枝剪", "浙江永康", "L4四星领航版 ¥128000/年"),
    ("水处理设备", "净水器、过滤器、反渗透设备、滤芯", "江苏无锡", "L3三星远航版 ¥69800/年"),
    ("化工助剂", "涂料助剂、塑料助剂、清洗剂、表面处理剂", "山东淄博", "L4四星领航版 ¥128000/年"),
]

STAGES = [
    ("外贸纯0-1起步", "无专职外贸人员，老板主动想试水，但担心踩坑和预算浪费", "新手", "L1"),
    ("初步试水阶段", "做过平台或展会，有零散询盘，但没有稳定成交闭环", "初级", "L1"),
    ("外贸成长瓶颈期", "有团队、有询盘，但获客成本高、转化不稳定、客户沉淀弱", "中级", "L2"),
    ("单一区域成熟期", "东南亚或中东存量稳定，但跨区域拓展认知不足", "中高级", "L2"),
    ("成熟团队全球化升级期", "内贸或外贸体量较大，想做多语种品牌和全域获客", "高级", "L3"),
]

BOSS_TRAITS = [
    "务实谨慎，重视投入产出，不喜欢空泛概念，愿意听本地同行案例。",
    "表达直接，喜欢看数据和落地路径，对服务承诺和合同边界敏感。",
    "愿意学习新东西，但害怕被平台年费和运营成本绑架。",
    "对销售专业度要求高，能接受高价方案，但要求看到阶段性成果。",
    "对外贸方向认可，但怕团队执行力跟不上，担心买完系统没人用。",
]

SCENARIO_NAMES = [
    "陌拜开场",
    "客户行业外贸趋势分析",
    "客户痛点诊断",
    "产品卖点介绍",
    "成单话术",
]

SALES_PRODUCTS = {
    "L1": "低风险试水出海方案、基础客户建联、轻量化官网和询盘承接",
    "L2": "独立站优化、多语种关键词布局、AI 客服承接、线索沉淀",
    "L3": "谷歌获客、社媒内容矩阵、询盘分级、业务员话术辅助",
    "L4": "多区域市场拓展、行业报告、海外品牌阵地、销售流程标准化",
    "L5": "全球化全域出海、品牌站群、多语种运营、AI 搜索卡位、组织级销售赋能",
}

PAIN_POOL = [
    "没有稳定海外获客渠道，客户来源靠熟人、展会或平台随机流量。",
    "网站和线上资产薄弱，海外采购商搜索不到企业，信任背书不足。",
    "销售线索分散在员工微信和表格里，客户资产不能沉淀到公司。",
    "报价、物流、交期、资质材料依赖人工整理，效率低且容易出错。",
    "平台同质化竞争严重，采购商只比价格，毛利被持续压缩。",
    "不懂海外区域差异，分不清哪些国家适合高端款、哪些适合走量款。",
    "外贸人员培养慢，新人不会谈单，老员工经验无法复制。",
    "老板担心投入后没有持续运营，最后变成一次性建站或摆设系统。",
]

HIDDEN_POOL = [
    "老板不是不想做外贸，而是怕自己不懂、团队不会、服务商只卖概念。",
    "客户真正想要的是低风险试错路径，希望有人把第一步拆清楚。",
    "客户表面压价，底层是在确认服务能不能落地、合同能不能保障。",
    "客户担心新增人力成本，希望用工具和托管服务替代盲目招人。",
    "客户已有存量业务，不想推翻原有模式，只想在可控范围内做增量。",
    "客户看重同行案例，只有看到同规模企业能跑通，才会放下戒备。",
]


def add_paragraph(document: Document, text: str = "", *, bold: bool = False) -> None:
    """向 Word 文档追加一段正文。

    这里统一字号，是为了生成的 50 份文档看起来一致，后续上传预览时也更清爽。
    """

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    """向 Word 文档追加标题。"""

    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "宋体"


def build_case_data(
    *,
    industry: str,
    products: str,
    city: str,
    stage: str,
    situation: str,
    level_code: str,
    trait: str,
    doc_index: int,
    scenario_index: int,
) -> dict[str, object]:
    """构造单个场景的客户背景、痛点和隐性心理。

    不同 doc_index 和 scenario_index 会错位选择素材，避免 50 份文档内容高度重复。
    """

    company_scale = [
        "12-18 人工贸小厂",
        "25-40 人成长型工厂",
        "60-90 人工贸结合企业",
        "120-180 人区域头部工厂",
        "220 人以上成熟制造企业",
    ][(doc_index + scenario_index) % 5]
    revenue = [
        "年营收约 600 万",
        "年营收约 1800 万",
        "年营收约 4200 万",
        "年营收约 1.2 亿",
        "年营收约 3 亿以上",
    ][(doc_index + scenario_index * 2) % 5]
    asset = [
        "无英文官网，仅有中文企业站",
        "有阿里国际站但运营粗放",
        "有模板独立站但没有谷歌排名",
        "有少量社媒账号但内容不稳定",
        "已有外贸团队但缺少多语种矩阵",
    ][(doc_index + scenario_index * 3) % 5]
    costs = [
        "平台年费和推广投入约 8 万",
        "展会和样册投入约 12 万",
        "线上获客成本约 18 万",
        "团队人工成本约 36 万",
        "年度市场预算超过 80 万",
    ][(doc_index + scenario_index) % 5]
    pain_a = PAIN_POOL[(doc_index + scenario_index) % len(PAIN_POOL)]
    pain_b = PAIN_POOL[(doc_index + scenario_index + 2) % len(PAIN_POOL)]
    pain_c = PAIN_POOL[(doc_index + scenario_index + 4) % len(PAIN_POOL)]
    hidden_a = HIDDEN_POOL[(doc_index + scenario_index) % len(HIDDEN_POOL)]
    hidden_b = HIDDEN_POOL[(doc_index + scenario_index + 3) % len(HIDDEN_POOL)]
    return {
        "asset": asset,
        "costs": costs,
        "pain_a": pain_a,
        "pain_b": pain_b,
        "pain_c": pain_c,
        "hidden_lines": [hidden_a, hidden_b, pain_b, pain_c],
        "case_lines": [
            f"企业：{city}{industry}制造工厂，主营{products}。",
            f"团队规模：{company_scale}，{revenue}。",
            f"外贸阶段：{stage}。",
            f"线上资产：{asset}。",
            f"年度投入现状：{costs}。",
            f"老板身份：{trait}",
            f"当前现状：{situation}；同时{pain_a}",
        ],
        "product_focus": SALES_PRODUCTS[level_code],
    }


def task_requirements(scenario: str) -> list[str]:
    """根据训练场景返回学员任务要求。"""

    if scenario == "陌拜开场":
        return [
            "请你模拟首次拜访客户时的开场白，要求如下：",
            "1、时间控制在30s-90s。",
            "2、语言自然，先降低推销感，再建立专业可信度。",
            "3、内容需包含个人介绍、公司介绍、客户现状判断、来意说明。",
            "4、最后预留一个能继续聊下去的钩子。",
        ]
    if scenario == "客户行业外贸趋势分析":
        return [
            "请围绕案例中客户所在行业，完成一段贴合客户行业的外贸趋势分析，要求如下：",
            "1、时间控制在1min-3min。",
            "2、至少包含行业趋势、客户当前阶段、机会区域、风险提醒。",
            "3、要结合客户规模和外贸阶段，不要泛泛讲大盘。",
            "4、最后自然引出一个适配产品包。",
        ]
    if scenario == "客户痛点诊断":
        return [
            "请结合案例中客户的实际情况，分析诊断当前客户痛点，要求如下：",
            "1、时间控制在1min-3min。",
            "2、痛点不少于4个，并且要贴合客户现状。",
            "3、表达要结构化，由显性问题递进到底层经营问题。",
            "4、能够给到初步适配方案或者套餐包。",
        ]
    if scenario == "产品卖点介绍":
        return [
            "请围绕案例中客户暴露出来的问题，针对性介绍 AI 盈出海的核心卖点和优势，要求如下：",
            "1、时间控制在1min-3min。",
            "2、至少解决客户当下3个痛点，并与产品卖点一一对应。",
            "3、表达过程中要抓住客户兴趣，引导客户愿意进一步看演示。",
            "4、不要空喊概念，要讲清楚怎么降低成本、提升转化、沉淀客户。",
        ]
    return [
        "请根据案例中客户的实际情况，完成一段成单推进话术，要求如下：",
        "1、时间控制在1min-3min。",
        "2、语气温和但有推进感，表达流畅，逻辑清晰。",
        "3、要打消客户顾虑，明确套餐价值、合同保障和下一步动作。",
        "4、至少给客户一个可落地执行的合作方案。",
    ]


def answer_text(
    *,
    scenario: str,
    industry: str,
    products: str,
    package: str,
    data: dict[str, object],
) -> list[str]:
    """生成单个场景的参考话术和答案。"""

    if scenario == "陌拜开场":
        return [
            f"匹配套餐：{package}",
            f"话术案例：老板您好，我是专门服务{industry}工厂做外贸增长和 AI 出海落地的。今天过来不是让您马上投入大项目，而是想结合咱们{products}这类产品的海外采购特点，先帮您判断现在有没有低风险试水的机会。",
            f"我最近接触了不少和您同规模的制造工厂，大家普遍不是不想出海，而是卡在没人懂外贸、怕平台投入高、怕招人没结果。您现在{data['asset']}，再加上{data['costs']}，如果继续靠零散方式摸索，试错成本会越来越高。",
            "我们能做的是先把海外客户从哪里来、怎么承接、怎么用 AI 辅助回复和沉淀线索这几件事拆清楚，不需要您一开始就扩团队。我想先用十分钟给您看一下同类工厂的出海路径，您判断有没有必要继续聊。",
        ]
    if scenario == "客户行业外贸趋势分析":
        return [
            f"适配套餐：{package}",
            f"话术案例：老板，{industry}这几年海外采购逻辑变化很明显。过去很多订单靠展会、平台和熟人转介绍，现在采购商越来越习惯先通过谷歌、AI 搜索、社媒内容和供应商官网做初筛。谁能被搜索到、谁能把工厂实力讲清楚，谁就更容易拿到第一轮询盘。",
            f"结合您现在的阶段，优势是产品线清楚，主营{products}，工厂有生产基础；短板是{data['asset']}，海外客户很难在第一时间判断您的实力。再加上{data['pain_a']}，这会直接影响询盘数量和信任转化。",
            f"建议不要一上来盲目砸广告，而是先用{package}把行业关键词、英文内容、AI 客服承接、线索沉淀和区域测试跑起来。这样既能保留现有业务，又能用较低风险验证海外增量市场。",
        ]
    if scenario == "客户痛点诊断":
        return [
            f"套餐适配：{package}",
            "核心显性卡点+隐性痛点：",
            f"1、渠道卡点：{data['pain_a']}",
            f"2、资产卡点：{data['asset']}，海外采购商即便听说过企业，也缺少可信的英文承接页面。",
            f"3、团队卡点：{data['pain_b']}",
            f"4、转化卡点：{data['pain_c']}",
            f"5、经营卡点：{data['costs']}，但没有形成可复用的获客和转化体系。",
            f"初步建议：先用{package}搭建可持续获客和客户承接闭环，把{data['product_focus']}跑通，再根据询盘质量决定是否扩大投放和团队配置。",
        ]
    if scenario == "产品卖点介绍":
        return [
            f"匹配产品包：{package}",
            f"话术原文：老板，您现在的问题不是产品不能卖海外，而是获客、承接、谈单、沉淀这几件事没有形成闭环。针对{industry}工厂，我们这套方案不是单纯建站，也不是只给您开账号，而是围绕线索进来以后怎么判断、怎么回复、怎么跟进、怎么沉淀客户来做。",
            f"第一，针对“{data['pain_a']}”，我们会做行业关键词和多语种内容布局，让海外采购商搜索{products}相关需求时，有机会看到您的工厂能力。",
            f"第二，针对“{data['pain_b']}”，AI 客服和话术辅助可以把常见询盘、报价、交期、认证问题标准化，新人也能按流程跟进。",
            f"第三，针对“{data['pain_c']}”，CRM 和客户沉淀机制能把客户资料留在公司，而不是散落在个人微信、邮箱和表格里。",
            f"所以{package}更适合您现在的状态，重点不是让您冒险扩张，而是用系统和服务把现有基础盘变成可复制的外贸增长流程。",
        ]
    return [
        f"套餐适配：{package}",
        f"成单话术：老板，前面我们把{industry}的海外机会、您工厂现在的短板和落地路径都聊清楚了。您现在最核心的顾虑，我理解不是不认可出海，而是担心投入后没人管、没有效果、合同里说不清楚。",
        f"这套{package}会把服务内容、交付节点、账号资产、内容产出、线索承接和复盘机制写清楚。我们不是让您单独招一支外贸团队去试错，而是先用一套标准化系统和陪跑服务，把{data['product_focus']}落到具体动作上。",
        f"从投入对比看，您现在{data['costs']}，但这些投入没有沉淀为稳定海外资产。把这笔预算的一部分转到可持续的线上获客和 AI 销售承接上，更容易形成长期复利。",
        "如果您认可，我们下一步可以先确认合同服务清单、交付周期和首批重点产品，然后安排资料收集和账号开通，先把第一阶段基础盘搭起来。",
    ]


def scoring_lines() -> list[str]:
    """返回评分模型可参考的通用评分点。"""

    return [
        "评分参考点：",
        "1、是否准确抓住客户当前阶段和真实顾虑。",
        "2、是否把显性痛点和隐性心理联系起来，而不是只罗列问题。",
        "3、是否能给出与案例匹配的产品包和下一步动作。",
        "4、表达是否自然、有层次、能形成继续沟通或成交推进。",
    ]


def build_document(doc_index: int) -> dict[str, object]:
    """生成一份 Word 文档，并返回目录行数据。"""

    industry, products, city, package = INDUSTRIES[(doc_index - 1) % len(INDUSTRIES)]
    stage, situation, trainee_level, level_code = STAGES[(doc_index - 1) % len(STAGES)]
    trait = BOSS_TRAITS[(doc_index - 1) % len(BOSS_TRAITS)]

    document = Document()
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"].font.size = Pt(10.5)

    title = f"LMS-销售训练案例扩展-{doc_index:02d}"
    heading = document.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(document, "资料定位：销售训练知识库扩展资料，用于 AI 客户角色生成、开放式训练对话和评分参考。")
    add_paragraph(document, f"行业主题：{industry}；地区样本：{city}；外贸阶段：{stage}；建议学员难度：{trainee_level}。")
    add_paragraph(document, "说明：本文档为模拟训练案例，结构保持“客户案例、任务要求、匹配答案、隐性心理、评分参考点”，方便上传后切片和检索。")

    for scenario_index, scenario in enumerate(SCENARIO_NAMES, 1):
        cn_num = ["一", "二", "三", "四", "五"][scenario_index - 1]
        add_heading(document, f"{cn_num}、{scenario}", level=1)
        add_paragraph(document, "客户案例：", bold=True)
        data = build_case_data(
            industry=industry,
            products=products,
            city=city,
            stage=stage,
            situation=situation,
            level_code=level_code,
            trait=trait,
            doc_index=doc_index,
            scenario_index=scenario_index,
        )
        for line in data["case_lines"]:
            add_paragraph(document, str(line))
        add_paragraph(document, "任务要求：", bold=True)
        for line in task_requirements(scenario):
            add_paragraph(document, line)
        add_paragraph(document, "匹配答案：", bold=True)
        for line in answer_text(scenario=scenario, industry=industry, products=products, package=package, data=data):
            add_paragraph(document, line)
        add_paragraph(document, "深层隐性心理（谈单核心突破口）：", bold=True)
        for line in data["hidden_lines"]:
            add_paragraph(document, str(line))
        add_paragraph(document, "补充细节：", bold=True)
        add_paragraph(
            document,
            f"客户更愿意听本行业、同规模、同阶段案例，不喜欢泛泛讲“趋势很好”。销售要把{industry}的产品特点、海外采购链路、风险控制和服务交付拆成具体动作。",
        )
        for line in scoring_lines():
            add_paragraph(document, line)

    filename = f"LMS-销售训练案例扩展-{doc_index:02d}-{industry}.docx"
    document.save(OUTPUT_DIR / filename)
    return {
        "序号": doc_index,
        "文件名": filename,
        "行业": industry,
        "地区样本": city,
        "外贸阶段": stage,
        "建议难度": trainee_level,
        "推荐套餐": package,
        "场景数量": 5,
    }


def write_manifest(rows: list[dict[str, object]]) -> None:
    """写出 CSV 目录，方便查看每份资料是什么行业和阶段。"""

    csv_path = OUTPUT_DIR / "案例目录.csv"
    with csv_path.open("w", newline="", encoding="utf-8-sig") as file:
        writer = csv.DictWriter(file, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_readme() -> None:
    """写出上传说明。"""

    readme = OUTPUT_DIR / "00-上传说明.md"
    readme.write_text(
        """# 销售训练案例扩展资料上传说明

本目录自动生成 50 份 LMS 销售训练案例扩展文档，每份文档包含 5 个训练场景，共 250 个训练场景。

建议上传方式：

1. 进入系统首页或销售陪练页面的训练资料上传入口。
2. 逐份上传 `LMS-销售训练案例扩展-xx-行业.docx`。
3. 上传后先查看质量报告和切片结构。
4. 确认无问题后点击“确认发布到训练库”。
5. 如果某份资料切片质量较低，可以点击“LLM 重新切分”。

文档结构：

- 客户案例
- 任务要求
- 匹配答案
- 深层隐性心理（谈单核心突破口）
- 补充细节
- 评分参考点

这些字段是为了让训练资料上传后更容易被切成客户画像、任务要求、标准话术、隐藏心理和评分规则。
""",
        encoding="utf-8",
    )


def main() -> None:
    """脚本入口。"""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = [build_document(index) for index in range(1, DOC_COUNT + 1)]
    write_manifest(rows)
    write_readme()
    print(f"生成目录: {OUTPUT_DIR}")
    print(f"生成 docx 数量: {len(list(OUTPUT_DIR.glob('*.docx')))}")
    print(f"目录文件: {OUTPUT_DIR / '案例目录.csv'}")
    print(f"说明文件: {OUTPUT_DIR / '00-上传说明.md'}")


if __name__ == "__main__":
    main()
